from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = "/Users/abdullah/Work/OS/mikeos-4.7.0/TOS_Beginners_Guide.docx"
BLUE = "235789"
DARK = "17324D"
LIGHT = "E8EEF5"
PALE = "F4F6F9"
GOLD = "B78018"
MUTED = "5B6573"
WHITE = "FFFFFF"

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.8)
sec.bottom_margin = Inches(0.75)
sec.left_margin = Inches(0.9)
sec.right_margin = Inches(0.9)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

def font(run, name="Aptos", size=None, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color: run.font.color.rgb = RGBColor.from_string(color)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.18

for name, size, color, before, after in [
    ("Title", 30, DARK, 0, 10), ("Subtitle", 14, MUTED, 0, 12),
    ("Heading 1", 18, BLUE, 16, 8), ("Heading 2", 14, BLUE, 12, 6),
    ("Heading 3", 11.5, DARK, 9, 4)]:
    s = styles[name]
    s.font.name = "Aptos Display" if "Heading" in name or name == "Title" else "Aptos"
    s._element.rPr.rFonts.set(qn("w:ascii"), s.font.name)
    s._element.rPr.rFonts.set(qn("w:hAnsi"), s.font.name)
    s.font.size = Pt(size)
    s.font.color.rgb = RGBColor.from_string(color)
    s.font.bold = name != "Subtitle"
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

for n in ["List Bullet", "List Number"]:
    s = styles[n]
    s.font.name = "Aptos"
    s.font.size = Pt(10.5)
    s.paragraph_format.left_indent = Inches(0.38)
    s.paragraph_format.first_line_indent = Inches(-0.19)
    s.paragraph_format.space_after = Pt(4)
    s.paragraph_format.line_spacing = 1.18

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd")) or OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    if shd.getparent() is None: tcPr.append(shd)

def margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc.get_or_add_tcPr()
    tcMar = tc.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar"); tc.append(tcMar)
    for k, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        x = tcMar.find(qn("w:"+k)) or OxmlElement("w:"+k)
        x.set(qn("w:w"), str(v)); x.set(qn("w:type"), "dxa")
        if x.getparent() is None: tcMar.append(x)

def set_cell_text(cell, text, bold=False, color=None, size=9.3, mono=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    font(r, "Menlo" if mono else "Aptos", size, bold, color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    margins(cell)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, h in enumerate(headers):
        set_cell_text(t.rows[0].cells[i], h, True, WHITE, 9.2)
        shade(t.rows[0].cells[i], BLUE)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val), False, None, 9.1)
            if len(t.rows) % 2 == 1: shade(cells[i], "F8FAFC")
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths): row.cells[i].width = Inches(w)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    return t

def p(text="", bold_start=None, italic=False, align=None):
    par = doc.add_paragraph()
    if align is not None: par.alignment = align
    if bold_start and text.startswith(bold_start):
        r = par.add_run(bold_start); font(r, bold=True)
        r = par.add_run(text[len(bold_start):]); font(r, italic=italic)
    else:
        r = par.add_run(text); font(r, italic=italic)
    return par

def bullet(text):
    par = doc.add_paragraph(style="List Bullet"); par.add_run(text); return par

current_num_id = None

def new_numbering_id():
    numbering = doc.part.numbering_part.element
    base_num_id = int(doc.styles["List Number"]._element.pPr.numPr.numId.val)
    base_num = next(n for n in numbering.findall(qn("w:num")) if int(n.get(qn("w:numId"))) == base_num_id)
    abstract_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    num_id = max(ids) + 1
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(num_id))
    aid = OxmlElement("w:abstractNumId"); aid.set(qn("w:val"), abstract_id); num.append(aid)
    override = OxmlElement("w:lvlOverride"); override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride"); start.set(qn("w:val"), "1"); override.append(start); num.append(override)
    numbering.append(num)
    return num_id

def number(text):
    global current_num_id
    if current_num_id is None: current_num_id = new_numbering_id()
    par = doc.add_paragraph(style="List Number")
    numPr = par._p.get_or_add_pPr().get_or_add_numPr()
    numPr.get_or_add_ilvl().val = 0
    numPr.get_or_add_numId().val = current_num_id
    par.add_run(text)
    return par

def code(lines):
    for line in lines.strip("\n").splitlines():
        par = doc.add_paragraph()
        par.paragraph_format.left_indent = Inches(0.25)
        par.paragraph_format.right_indent = Inches(0.15)
        par.paragraph_format.space_after = Pt(0)
        par.paragraph_format.line_spacing = 1.0
        shade_par(par, "F2F4F7")
        r = par.add_run(line); font(r, "Menlo", 8.6, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def shade_par(par, fill):
    pPr = par._p.get_or_add_pPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); pPr.append(shd)

def callout(label, text, fill=PALE):
    t = doc.add_table(rows=1, cols=1); t.autofit = False; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0,0); shade(c, fill); margins(c, 130, 180, 130, 180)
    c.text = ""; par = c.paragraphs[0]; par.paragraph_format.space_after = Pt(0)
    r = par.add_run(label + "  "); font(r, bold=True, color=BLUE)
    r = par.add_run(text); font(r, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def heading(text, level=1):
    global current_num_id
    current_num_id = None
    doc.add_heading(text, level=level)

def page_break(): doc.add_page_break()

def add_page_field(par):
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); par._p.append(fld)

# Header/footer
hp = sec.header.paragraphs[0]
hp.text = "TOS 4.7.0  |  Beginner's Operating System Guide"
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(hp.runs[0], size=8.5, color=MUTED)
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run("TOS Study Guide  •  Page "); font(r, size=8.5, color=MUTED); add_page_field(fp)

# Cover
doc.add_paragraph().paragraph_format.space_after = Pt(90)
kick = doc.add_paragraph(); kick.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = kick.add_run("CLASS PRESENTATION EDITION"); font(r, size=10, bold=True, color=GOLD)
title = doc.add_paragraph(style="Title"); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("How TOS Works")
sub = doc.add_paragraph(style="Subtitle"); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("A beginner-friendly guide to booting, the kernel, BIOS services, keyboard input, FAT12 files, and the command line")
doc.add_paragraph().paragraph_format.space_after = Pt(45)
callout("ONE-SENTENCE DESCRIPTION", "TOS is a small, single-tasking, 16-bit x86 operating system that runs in BIOS real mode and uses a FAT12 floppy filesystem.", "E8EEF5")
p("Prepared for Abdullah • Based on the TOS 4.7.0 source tree • 11 July 2026", align=WD_ALIGN_PARAGRAPH.CENTER)
p("Scope: kernel and CLI fundamentals. Bundled programs and applications are intentionally excluded.", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
page_break()

heading("How to use this guide", 1)
p("Read Chapters 1–4 first to build the mental model. Chapters 5–8 explain the hardware-facing services and command line. Chapter 9 gives command-by-command behavior. The final sections contain a presentation script, likely questions, and a glossary.")
heading("Learning goals", 2)
for x in [
    "Explain what an operating system does and where TOS fits.",
    "Describe the complete path from power-on to the TOS prompt.",
    "Explain real mode, segments, registers, memory layout, interrupts, and the stack at a beginner level.",
    "Trace a keypress from the keyboard through BIOS to the CLI input buffer.",
    "Explain how FAT12 stores files and how DIR discovers filenames.",
    "Trace how the CLI parses, recognizes, and runs commands.",
    "Discuss TOS's strengths and limitations honestly during a class presentation."]:
    bullet(x)

heading("Contents", 1)
for x in ["1. OS fundamentals", "2. Architecture at a glance", "3. From power-on to kernel", "4. CPU and memory model", "5. Screen and keyboard", "6. FAT12 and disk I/O", "7. The CLI loop", "8. Why DIR works", "9. Built-in command reference", "10. Build and run workflow", "11. Limitations and comparisons", "12. Presentation script and Q&A", "Glossary and source map"]:
    p(x)
page_break()

heading("1. Operating-system fundamentals", 1)
heading("What is an operating system?", 2)
p("An operating system (OS) is the trusted software layer between hardware and the user or other software. It decides how the CPU starts, how memory is arranged, how input and output happen, how data is stored, and what interface the user receives. Modern systems such as macOS and Linux do this with millions of lines of code. TOS intentionally demonstrates the same ideas with a much smaller design.")
heading("What TOS is—and is not", 2)
table(["Property", "TOS"], [
    ("CPU/platform", "Intel-compatible x86; 386 or newer CPU instructions, but 16-bit real mode"),
    ("Execution", "Single task at a time; no scheduler or processes"),
    ("Memory", "One 64 KiB segment shared by kernel and loaded code"),
    ("Hardware access", "Primarily BIOS interrupts"),
    ("Filesystem", "FAT12 root directory on a 1.44 MB floppy-style image"),
    ("Interface", "Text UI with a program selector and a command line"),
    ("Protection", "No users, permissions, privilege rings, or memory isolation"),
    ("Goal", "Education, clarity, and small size—not production security or performance"),
], [1.4, 5.0])
callout("KEY IDEA", "TOS is still an operating system because it boots directly, initializes its environment, exposes reusable kernel services, handles input/output, manages files, and provides a user interface—even though it is intentionally tiny.")

heading("Kernel, bootloader, shell, and BIOS", 2)
table(["Term", "Meaning in TOS"], [
    ("BIOS", "Firmware interface already present when the machine starts; offers interrupts for screen, keyboard, disk, clock, and reboot."),
    ("Bootloader", "The first 512-byte sector. It finds KERNEL.BIN on FAT12, loads it, and jumps to it."),
    ("Kernel", "The core binary containing startup logic and reusable OS routines."),
    ("CLI/shell", "A loop inside the kernel that reads text, recognizes commands, and calls kernel routines."),
    ("System-call vectors", "A fixed jump table at the start of the kernel that gives stable entry addresses to OS services."),
], [1.4, 5.0])
heading("2. Architecture at a glance", 1)
heading("The whole system in one flow", 2)
table(["Stage", "What happens", "Main mechanism"], [
    ("1. Firmware", "BIOS selects the boot device and loads sector 0 at physical address 0x7C00.", "BIOS boot convention"),
    ("2. Boot sector", "The 512-byte loader reads the FAT12 root directory and searches for KERNEL.BIN.", "INT 13h disk reads"),
    ("3. File loading", "The loader follows the file's FAT12 cluster chain and copies its sectors to segment 0x2000.", "FAT12 + CHS conversion"),
    ("4. Kernel entry", "A far jump transfers execution to 0x2000:0000, whose first vector jumps to os_main.", "x86 jump table"),
    ("5. Initialization", "The kernel sets stack and segments, reads drive geometry, configures text behavior, and seeds randomness.", "Registers + BIOS"),
    ("6. Interface", "The welcome dialog chooses a selector or the CLI. CLI repeats prompt → input → parse → dispatch.", "Kernel routines"),
], [0.8, 3.7, 1.9])

heading("The BIOS interrupt bridge", 2)
p("In real mode, the kernel can call firmware routines using the x86 INT instruction. AH usually selects a sub-function; other registers carry inputs and outputs. The carry flag often reports success or failure.")
table(["Interrupt", "Service", "How TOS uses it"], [
    ("INT 10h", "Video", "Print characters, clear regions, move/read the cursor, set colors."),
    ("INT 13h", "Disk", "Reset the drive and read/write floppy sectors."),
    ("INT 16h", "Keyboard", "Check for a key and retrieve ASCII plus scan code."),
    ("INT 1Ah", "Clock", "Read time/date and timer ticks."),
    ("INT 19h", "Bootstrap", "Reboot by asking BIOS to restart the boot sequence."),
], [1.0, 1.2, 4.2])
callout("WHY THIS MATTERS", "TOS does not need complex hardware drivers for its basic devices. BIOS already knows the hardware, so the kernel asks BIOS to do the low-level work. The tradeoff is that BIOS services are old, slow, and limited.")

heading("3. From power-on to the TOS kernel", 1)
heading("Step 1: BIOS loads the boot sector", 2)
p("A PC firmware does not understand the TOS kernel format. It follows a standard rule: read the first 512-byte sector of the selected boot device into memory at physical address 0x7C00 and execute it. The last two bytes must be the boot signature 0x55AA. TOS's bootloader must therefore fit in exactly one sector.")
heading("Step 2: establish a safe stack and data segment", 2)
p("The loader sets DS so labels refer to its own data and creates a stack above its buffer. CLI, STI, SS, and SP are involved: interrupts are briefly disabled while the stack changes, SS selects the stack segment, and SP selects the top offset.")
heading("Step 3: locate KERNEL.BIN", 2)
p("The loader knows the 1.44 MB FAT12 geometry. It reads logical sectors 19–32, which hold 224 root-directory entries. Each entry is 32 bytes. It compares the first 11 bytes against the FAT short-name form KERNEL␠␠BIN.")
heading("Step 4: follow the FAT12 cluster chain", 2)
p("The matching directory entry contains the first cluster number. FAT12 stores the next-cluster value in 12 bits, so the loader performs special odd/even calculations to extract each entry. With one 512-byte sector per cluster, it reads each cluster into memory and advances until the end-of-chain marker.")
heading("Step 5: jump into the kernel", 2)
p("The kernel is loaded at segment 0x2000. In real mode that segment begins at physical address 0x20000 because physical = segment × 16 + offset. Execution enters offset 0, where the system-call vector table's first instruction jumps to os_main.")
code("""BIOS → loads boot sector at 0000:7C00
boot sector → searches FAT12 root for KERNEL.BIN
boot sector → loads file at 2000:0000
2000:0000 → jmp os_main
os_main → initializes TOS and opens the interface""")
page_break()

heading("4. CPU and memory model", 1)
heading("Real mode in plain English", 2)
p("Real mode is the x86 startup mode inherited from the original 8086. Addresses are formed from a 16-bit segment and 16-bit offset. There is no memory protection: code can read or overwrite almost anything. TOS uses 386 instructions where convenient, but deliberately stays in real mode so BIOS calls and segmentation remain simple.")
heading("Registers you should recognize", 2)
table(["Register", "Beginner meaning", "Examples in TOS"], [
    ("AX, BX, CX, DX", "General working registers, each also split into high/low bytes.", "AH selects BIOS functions; AL carries characters; CX counts; DX carries cursor or disk values."),
    ("SI / DI", "Source and destination indexes.", "Walk through strings, compare names, copy buffers."),
    ("SP / SS", "Stack pointer and stack segment.", "CALL, RET, PUSH, and POP depend on them."),
    ("CS / IP", "Current code segment and instruction offset.", "Together identify the executing instruction."),
    ("DS / ES", "Data segments.", "TOS points them at 0x2000 so code and data share one segment."),
    ("Flags", "One-bit results and controls.", "Carry reports disk/file errors; zero reports comparisons or missing keys."),
], [1.1, 2.3, 3.0])

heading("TOS memory map", 2)
table(["Offset in segment 0x2000", "Purpose"], [
    ("0x0000 onward", "Kernel jump vectors, kernel code, and kernel data"),
    ("0x6000 (24 KiB)", "8 KiB disk_buffer used for filesystem operations"),
    ("0x8000 (32 KiB)", "Load address for a file or externally executed code"),
    ("SS=0, SP near 0xFFFF", "Kernel stack below physical 0x10000; it grows downward and is outside segment 0x2000"),
], [2.0, 4.4])
p("This layout is compact but fragile. Code and most data share segment 0x2000, while the kernel stack uses segment 0. There is no MMU isolation, and a bad pointer or runaway stack can corrupt the system.")

heading("System-call vectors", 2)
p("The kernel begins with three-byte JMP instructions at fixed offsets: print string at 0x0003, clear screen at 0x0009, wait for key at 0x0012, load file at 0x0021, and so on. Internal routines may move during assembly, but these entry slots remain stable. This is a tiny application binary interface (ABI).")
callout("CLASS ANALOGY", "Think of the vector table as a reception desk. Callers know the desk number, and the desk forwards them to the current office. The office may move; the public desk number must not.")

heading("5. Screen and keyboard handling", 1)
heading("How text reaches the screen", 2)
p("Strings are stored as bytes ending in zero. os_print_string loads one byte at a time with LODSB. If it is zero, the routine returns. Otherwise AH=0x0E selects BIOS teletype output and INT 10h prints the character at the cursor, advancing it automatically.")
code("""SI → address of text
repeat:
    AL = next byte
    if AL == 0: return
    AH = 0x0E
    INT 0x10
    repeat""")
p("The screen is classic 80-column × 25-row text mode. os_clear_screen uses INT 10h function 6 to scroll/clear the full rectangle, while os_move_cursor uses function 2 with DH=row and DL=column.")

heading("How a keypress reaches the CLI", 2)
number("The physical keyboard/controller reports a key event to the BIOS.")
number("BIOS translates it and places a key record in its keyboard buffer.")
number("os_wait_for_key calls INT 16h function 0x11 to check the enhanced keyboard buffer.")
number("If empty, the CPU executes HLT to save work and wakes on an interrupt, then checks again.")
number("When a key exists, INT 16h function 0x10 removes it and returns AX: AL is usually ASCII; AH is the scan code.")
number("os_input_string interprets editing keys, echoes characters, and stores the resulting zero-terminated line in the CLI input buffer.")

heading("Blocking versus non-blocking input", 2)
table(["Routine", "Behavior", "Use"], [
    ("os_wait_for_key", "Waits until a key exists; uses HLT while idle.", "Menus, prompts, and any action that must wait."),
    ("os_check_for_key", "Checks once and returns AX=0 when no key exists.", "Loops that must continue doing work without pausing."),
], [1.5, 2.7, 2.2])
callout("IMPORTANT DISTINCTION", "The kernel is not directly decoding USB packets or electrical scan codes. QEMU emulates a PC keyboard; its virtual BIOS performs that lower-level translation; TOS consumes the BIOS result through INT 16h.")
page_break()

heading("6. FAT12 and disk I/O", 1)
heading("What FAT12 contributes", 2)
p("FAT12 is a simple filesystem common on floppy disks. It divides storage into sectors and clusters, stores file metadata in fixed-size directory entries, and keeps a File Allocation Table (FAT) describing which cluster follows which.")
table(["Region", "Logical sectors", "Purpose"], [
    ("Boot sector", "0", "Boot code plus BIOS Parameter Block describing the disk"),
    ("FAT copy 1", "1–9", "Primary 12-bit cluster chain table"),
    ("FAT copy 2", "10–18", "Backup copy"),
    ("Root directory", "19–32", "224 entries × 32 bytes"),
    ("Data area", "33 onward", "Actual contents of files"),
], [1.5, 1.4, 3.5])

heading("The 8.3 filename rule", 2)
p("A normal FAT12 short-name entry stores eight bytes for the base name and three for the extension. Names are uppercase and space-padded on disk—for example, README.TXT becomes README␠␠TXT. TOS converts between this 11-byte internal form and the dotted form shown to users.")
heading("Logical sector to physical CHS", 2)
p("BIOS INT 13h expects cylinder/head/sector values for this floppy interface. disk_convert_l2hts divides a logical sector number by sectors-per-track and number of sides to calculate sector, head, and track. Sector numbering begins at 1, not 0.")
heading("Reading a file", 2)
number("Uppercase and convert the requested filename to 8.3 format.")
number("Reset the drive, then read the 14 root-directory sectors into disk_buffer.")
number("Scan 32-byte entries, skipping deleted, long-name, directory, and volume-label entries.")
number("When the name matches, read its first cluster and byte size.")
number("Read the FAT, follow the cluster chain, and copy each sector to the caller's requested RAM address.")
number("Return the file size in EBX; set the carry flag if it fails.")

heading("7. The command-line loop", 1)
heading("A shell is a loop, not magic", 2)
p("TOS's CLI is a routine inside the kernel. It prints a version and help line once, then repeats the same small algorithm until EXIT returns to the welcome screen.")
code("""clear screen
print version and command list
loop:
    clear command buffer
    print '> '
    read up to 64 characters
    trim and tokenize input
    uppercase the command word
    compare against built-in command names
    run the matching handler
    otherwise try .BIN, then .BAS
    if nothing matches, print an error
    go back to loop""")

heading("Input buffers and parsing", 2)
table(["Item", "Size/role"], [
    ("input", "64-byte buffer holding the line typed by the user"),
    ("command", "32-byte buffer holding the first word / command name"),
    ("param_list", "Pointer to the remaining words after tokenization"),
    ("dirlist", "1024-byte buffer for comma-separated filenames"),
], [1.5, 4.9])
p("The CLI removes trailing spaces, ignores an empty line, and uses the first space as the split point. It copies the first token into command and preserves a pointer to the rest. It then uppercases the command token, which is why dir, Dir, and DIR all work.")

heading("Command dispatch", 2)
p("Dispatch is a sequence of string comparisons, not a table or separate process. For each known command, the kernel points SI at the typed word and DI at a constant such as 'DIR', then calls os_string_compare. When the strings match, the routine sets carry, and a conditional jump transfers control to that command handler.")
code("""compare input with 'EXIT' → if equal, jump to exit
compare input with 'HELP' → if equal, jump to print_help
compare input with 'CLS'  → if equal, jump to clear_screen
compare input with 'DIR'  → if equal, jump to list_directory
...""")
callout("WHY UNKNOWN COMMANDS FAIL", "Typing cd, :q, or q produces 'No such command or program' because none is a built-in comparison and no matching CD.BIN/CD.BAS (or equivalent) file exists.")

heading("8. Why DIR works: complete walkthrough", 1)
number("You type dir and press Enter. BIOS places each keystroke in its keyboard buffer.")
number("os_input_string retrieves and echoes the characters, producing the zero-terminated string 'dir'.")
number("The CLI trims spaces, tokenizes the line, copies the first word, and uppercases input to 'DIR'.")
number("os_string_compare matches it against dir_string. The carry flag signals equality, so execution jumps to list_directory.")
number("list_directory sets AX to the 1024-byte dirlist buffer and calls os_get_file_list.")
number("os_get_file_list reads the FAT12 root directory into the kernel's disk_buffer using BIOS INT 13h.")
number("It walks 32-byte entries, ignoring deleted entries, long-filename markers, directories, and the volume label.")
number("For every valid short filename, it removes padding, inserts a dot before a real extension, and appends the result to a comma-separated string.")
number("Back in list_directory, the CLI reads that string one character at a time. Commas mean 'next filename'.")
number("It calculates columns 0, 20, 40, and 60, moves the cursor with BIOS INT 10h, and prints four filenames per row.")
number("At the zero terminator, it prints a newline and returns to the prompt loop.")
callout("THE BIG LESSON", "DIR is a pipeline across nearly every OS layer: keyboard → input buffer → parser → command dispatcher → filesystem parser → BIOS disk service → formatted screen output.", "FFF4D6")
heading("9. Built-in command reference", 1)
p("These are kernel CLI commands, not bundled applications. Parameters are separated by spaces; FAT short-name limitations apply.")
table(["Command", "Purpose", "Internal path"], [
    ("DIR", "List filenames in four columns.", "os_get_file_list → format commas into columns"),
    ("LS", "Detailed FAT directory listing with metadata and paging.", "Reads root entries and formats fields"),
    ("CAT file", "Print a text file.", "os_file_exists → os_load_file at 0x8000 → INT 10h per byte"),
    ("SIZE file", "Show byte size.", "os_get_file_size → integer-to-string"),
    ("COPY src dst", "Duplicate a file.", "load src at 0x8000 → os_write_file dst"),
    ("REN old new", "Rename a file.", "check destination → os_rename_file"),
    ("DEL file", "Delete a file.", "os_remove_file marks entry and frees clusters"),
    ("CLS", "Clear screen.", "os_clear_screen / INT 10h"),
    ("HELP", "Print command list.", "os_print_string"),
    ("VER", "Print TOS version.", "os_print_string"),
    ("TIME", "Show BIOS clock time.", "os_get_time_string / INT 1Ah"),
    ("DATE", "Show BIOS clock date.", "os_get_date_string / INT 1Ah"),
    ("EXIT", "Leave CLI for welcome screen.", "RET from os_command_line"),
], [1.0, 2.4, 3.0])

heading("Command examples", 2)
code("""> DIR
> CAT README.TXT
> SIZE KERNEL.BIN
> COPY README.TXT NOTES.TXT
> REN NOTES.TXT CLASS.TXT
> DEL CLASS.TXT
> TIME
> VER
> EXIT""")
heading("Error behavior", 2)
for x in [
    "Missing parameter: 'No filename or not enough filenames'.",
    "Missing file: 'File not found'.",
    "Existing destination: 'Target file already exists!'.",
    "Write problem: write-protected disk, invalid name, or no available space/cluster.",
    "Unknown token: CLI tries executable extensions before printing 'No such command or program'."]:
    bullet(x)

heading("10. Build and run workflow", 1)
heading("What the build script does", 2)
number("NASM assembles bootload.asm into exactly 512 bytes of machine code.")
number("NASM assembles kernel.asm; INCLUDE directives pull feature files into one kernel binary.")
number("The script assembles other binaries, copies a base floppy image, and writes the boot sector over sector 0.")
number("macOS attaches the FAT image, then the script copies KERNEL.BIN and data onto it.")
number("mkisofs wraps the bootable floppy image into an El Torito bootable ISO.")
code("""cd /Users/abdullah/Work/OS/mikeos-4.7.0
sudo ./build-macos.sh

qemu-system-i386 -m 16M -boot d \\
  -cdrom disk_images/mikeos.iso \\
  -display cocoa,full-screen=on,zoom-to-fit=on \\
  -no-reboot""")
heading("Why QEMU works", 2)
p("QEMU emulates the x86 CPU, BIOS, VGA text display, keyboard controller, floppy/CD boot path, and other PC hardware. TOS believes it is running on a physical legacy PC. Your macOS keyboard and window are translated into the virtual hardware that TOS expects.")

heading("11. Limitations and comparison with modern OSes", 1)
table(["Topic", "TOS", "Modern OS"], [
    ("CPU mode", "16-bit real mode", "64-bit protected/long mode"),
    ("Processes", "One flow at a time", "Preemptive multitasking and scheduling"),
    ("Memory safety", "None", "Virtual memory, page permissions, isolation"),
    ("Drivers", "Mostly BIOS", "Native drivers, plug-and-play, power management"),
    ("Filesystem", "FAT12 root only", "Directories, journaling, permissions, large storage"),
    ("Security", "No accounts or privilege boundary", "Users, ACLs, sandboxing, secure boot"),
    ("Networking", "No kernel network stack", "TCP/IP, Wi-Fi, firewalls"),
    ("Scale", "Tiny and readable", "Millions of lines and many subsystems"),
], [1.3, 2.4, 2.7])
heading("What TOS teaches exceptionally well", 2)
for x in ["The boot contract between firmware and software.", "How machine code, registers, flags, and memory addresses implement abstractions.", "Why a filesystem is a data structure on sectors—not magic.", "How input, parsing, dispatch, and output form a shell.", "How stable function entry points become an API.", "Why modern protection and abstraction layers were invented."]:
    bullet(x)

heading("12. Class presentation plan", 1)
heading("A 7–10 minute talk", 2)
table(["Time", "What to say/show"], [
    ("0:00–0:45", "Define TOS: tiny 16-bit x86 educational OS, BIOS real mode, FAT12."),
    ("0:45–2:00", "Show boot chain: BIOS → boot sector → FAT12 → KERNEL.BIN → os_main."),
    ("2:00–3:15", "Explain BIOS interrupts: 10h screen, 13h disk, 16h keyboard, 1Ah clock."),
    ("3:15–4:15", "Explain shared 64 KiB memory map and fixed system-call vectors."),
    ("4:15–6:30", "Live demo CLI. Run VER, TIME, DIR, CAT README.TXT, SIZE KERNEL.BIN."),
    ("6:30–8:00", "Trace DIR from keypress to FAT directory and back to display."),
    ("8:00–9:00", "Contrast TOS with modern OS: no multitasking, protection, users, or networking."),
    ("9:00–10:00", "Conclude: small enough to trace, complete enough to teach core layers."),
], [1.0, 5.4])

heading("Suggested opening", 2)
p("“Today I am demonstrating TOS, a small 16-bit operating system for x86 PCs. It boots directly from a FAT12 disk image, loads its own kernel, communicates with virtual hardware through BIOS interrupts, and provides a command-line interface. Its limitations are deliberate: because it is small, we can trace a keypress or a DIR command through the whole system.”", italic=True)
heading("Live-demo sequence", 2)
code("""VER
TIME
DATE
DIR
CAT README.TXT
SIZE KERNEL.BIN
CLS
HELP
EXIT""")
callout("DEMO SAFETY", "Use read-only commands during class. Avoid DEL, REN, and COPY unless you have rebuilt the image or made a backup. Keep the QEMU command ready in a terminal.", "FFF4D6")

heading("Questions your teacher may ask", 2)
table(["Question", "Strong short answer"], [
    ("Is this really an OS?", "Yes. It boots without a host OS inside the VM, initializes the machine environment, manages files and I/O, exposes kernel services, and provides a UI."),
    ("Why 16-bit real mode?", "It keeps the architecture small and lets the kernel call BIOS services directly. The cost is weak memory capacity and no protection."),
    ("Where is the kernel loaded?", "At segment 0x2000, physical address 0x20000. Offset 0 contains its stable jump vectors."),
    ("How is keyboard input handled?", "QEMU emulates the device, BIOS buffers and translates keys, and TOS calls INT 16h to check/read them."),
    ("Why does DIR work?", "The CLI recognizes the string, reads FAT12 root entries through INT 13h, formats valid 8.3 names, and prints them through INT 10h."),
    ("Why no CD command?", "The filesystem model only scans a flat FAT12 root directory; subdirectory navigation is not implemented."),
    ("Does it multitask?", "No. One kernel routine or loaded task runs until it returns. There is no scheduler or process isolation."),
    ("What did you customize?", "The visible product branding was changed to TOS, with the welcome greeting 'Welcome Abdullah'; internal compatibility names were preserved."),
], [2.0, 4.4])
page_break()

heading("Glossary", 1)
terms = [
    ("ABI", "Rules and stable binary entry points that allow compiled code to call services."),
    ("Assembly", "Human-readable notation for CPU instructions."),
    ("BIOS", "Legacy PC firmware and its callable hardware services."),
    ("Boot sector", "The first 512-byte sector executed by BIOS."),
    ("Buffer", "A reserved RAM area used to temporarily hold data."),
    ("Carry flag", "A CPU status bit used here to report comparisons and errors."),
    ("CHS", "Cylinder/head/sector addressing used by BIOS floppy services."),
    ("CLI", "Command-line interface: text prompt, parser, and command dispatcher."),
    ("Cluster", "Filesystem allocation unit; one sector in this FAT12 image."),
    ("FAT12", "Filesystem with 12-bit entries linking file clusters."),
    ("Interrupt", "A CPU mechanism that transfers control to a handler; BIOS services use software INT instructions."),
    ("Kernel", "The trusted core providing startup and OS services."),
    ("Real mode", "Original x86 mode with segmented addressing and no memory protection."),
    ("Register", "Small, fast storage location inside the CPU."),
    ("Segment:offset", "Real-mode address pair; physical address = segment × 16 + offset."),
    ("Shell", "User interface that reads commands and dispatches work."),
    ("Stack", "Last-in-first-out memory used by CALL/RET and PUSH/POP."),
    ("System call", "A defined entry into an OS service."),
    ("Zero-terminated string", "Characters followed by byte 0, marking the end."),
]
table(["Term", "Definition"], terms, [1.5, 4.9])

heading("Source map for further study", 1)
table(["File", "What to study"], [
    ("source/bootload/bootload.asm", "Boot sector, FAT12 kernel search, cluster loading, BIOS disk calls"),
    ("source/kernel.asm", "Jump vectors, memory constants, os_main, welcome/interface flow"),
    ("source/features/cli.asm", "Prompt loop, parsing, dispatch, built-in commands"),
    ("source/features/keyboard.asm", "Blocking and non-blocking INT 16h input"),
    ("source/features/screen.asm", "INT 10h printing, cursor, clearing, dialogs"),
    ("source/features/disk.asm", "FAT12 listing, loading, writing, deleting, CHS conversion"),
    ("source/features/string.asm", "String compare, uppercase, tokenize, parse"),
    ("build-macos.sh", "Assembly, disk-image population, and ISO creation"),
], [2.6, 3.8])

heading("Final mental model", 1)
callout("REMEMBER THIS CHAIN", "Firmware loads the boot sector → the bootloader understands enough FAT12 to load the kernel → the kernel uses BIOS services and shared memory to implement reusable routines → the CLI turns typed strings into calls to those routines → results return to the screen.", "E8EEF5")
p("If you can explain that chain, trace DIR, and name the major limitations, you understand the most important ideas in this operating system.")

# Keep tables readable across pages and apply fixed cell padding.
for t in doc.tables:
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for c in row.cells: margins(c)
    trPr = t.rows[0]._tr.get_or_add_trPr()
    rep = OxmlElement("w:tblHeader"); rep.set(qn("w:val"), "true"); trPr.append(rep)

doc.core_properties.title = "How TOS Works: A Beginner's Operating System Guide"
doc.core_properties.subject = "TOS 4.7.0 architecture, boot process, keyboard, FAT12, and CLI"
doc.core_properties.author = "Abdullah"
doc.core_properties.keywords = "TOS, operating system, x86, BIOS, FAT12, assembly, CLI"
doc.save(OUT)
print(OUT)
